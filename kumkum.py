#!/usr/bin/env python
from openpyxl import load_workbook
from openpyxl import Workbook
#from openpyxl.styles import Border, Side, PatternFill, Font, GradientFill, Alignment, Color, Fill
from openpyxl import formatting, styles
from openpyxl.cell import Cell
import string
import sys

class KumKum(object):
	def __init__(self):
		self.main()

	def showFilesOnDir(self, path = 'd:/kuman'):
		import os
		data = []
		number = 1
		print("============ DAFTAR FILE ============")
		print("=====================================")
		for file in os.listdir(path):
			if file.endswith(".xlsx") and "~$" not in file:
				if file not in data:
					print(str(number) + '.', file)
					data.append({'id': number, 'name': file})
					number += 1
		return data
		
	def pickAFile(self, files):
		data = {}
		if len(files) > 0:
			# Pick a file
			pickXlsx = input("Pilih salah satu: ")
	
			for d in files:
				if d['id'] == int(pickXlsx):
					data = d
		return data
		
	def pickMonthYear(self, months):
		data = {"month": 5, "year": 2017}
		print("===== FILE TERSEBUT ADALAH BULAN =====")
		print("======================================")
		for key, value in months.items():
			print(value['alias'], value['name'])
		pickMonth = input("Pilih salah satu: ")
		pickYear = input("Masukkan tahun: ")
		data['month'] = int(pickMonth)
		data['year'] = int(pickYear)
		return data
		
	def loadDictionary(self, path = 'd:/kuman', file = 'dictionary.json'):
		import json
		data = []
		with open(file, 'r') as f:
			data = json.load(f)
		return data
	
	def openFile(self, path = 'd:/kuman', file = 'mei.xlsx'):
		data = None
		filename = path + '/' + file
		#data = load_workbook(filename=filename, data_only=True, read_only=True)
		data = load_workbook(filename=filename, data_only=True)
		return data
		
	def getSheets(self, file):
		data = []
		if file:
			for sheet_name in file.sheetnames:
				data.append(sheet_name)
		return data
	
	def sheetCheckIfNotBlank(self, ws = None):
		data = True
		#ws = file[sheetName]
		if ws['A3'].value:
			if ws['A3'].value.lower() == 'tidak ada':
				data = False
		return data
		
	def sheetGetBacteries(self, ws = None):
		data = {}
		
		# Row contain "organisma"
		wb_row_organisma_start = 0
		wb_row_organisma_end = 0
		wb_row_organisma = 0
		
		# Find row contain "organisma" word
		for wb_row_index in range(1, 100):
			if ws['B' + str(wb_row_index)].value:
				if ws['B' + str(wb_row_index)].value.lower() == 'organisma':
					wb_row_organisma_start = wb_row_index
					wb_row_organisma = wb_row_index
				# Find organisma lowest, by find "nama antibiotik" word
				elif ws['B' + str(wb_row_index)].value.lower() == 'nama antibiotik':
					if wb_row_organisma_end == 0 or wb_row_organisma_end > wb_row_index:
						wb_row_organisma_end = wb_row_index
			
		wb_row_organisma_start += 1
		wb_row_organisma_end -= 1
		# Create variable to detect last line of organisma
		wb_row_organisma_end_temp = wb_row_organisma_start
		# Iterate A and B column for organisma
		if wb_row_organisma_start > 0:
			for wb_row_index in range(wb_row_organisma_start, wb_row_organisma_end):
				if ws['A' + str(wb_row_index)].value and ws['B' + str(wb_row_index)].value:
					wb_row_organisma_end_temp = wb_row_index
					data[ws['A' + str(wb_row_index)].value.lower()] = {"alias": ws['A' + str(wb_row_index)].value.lower(), "name": ws['B' + str(wb_row_index)].value.lower(), "bactery_row": wb_row_index, "antibiotic_row": 0, "speciments": {"total_value": 0}, "antibiotics": {"r": {"max_value": float(0), "antibiotics_row": 0, "antibiotics": {}}, "i": {"max_value": float(0), "antibiotics_row": 0, "antibiotics": {}}, "s": {"max_value": float(0), "antibiotics_row": 0, "antibiotics": {}}}}
					#data.append({"alias": ws['A' + str(wb_row_index)].value.lower(), "name": ws['B' + str(wb_row_index)].value.lower(), "bactery_row": wb_row_index, "speciments": {}})
						
		wb_row_organisma_end = wb_row_organisma_end_temp
		return {"data": data, "start_row": wb_row_organisma_start, "end_row": wb_row_organisma_end, "organisma_row": wb_row_organisma}
	
	def sheetAppendBacteriesWithSpeciments(self, ws = None, bacteries = [], start_row = 0, end_row = 0, organisma_row = 0, dictionary_speciments = {}):
		data = bacteries
		# Get speciment used
		# coloumn used from F to Z
		# Get which column used
		wb_column_organisma_end = 6 # Column F
		specimentColumns = []
		for cell in ws.iter_cols(min_col=6, max_col=26, min_row=organisma_row, max_row=organisma_row):
			for c in cell:
				if c.value:
					specimentColumns.append({"speciment": c.value, "column_position": wb_column_organisma_end})
					wb_column_organisma_end += 1
		
		if wb_column_organisma_end > 6:
			wb_column_organisma_end -= 1
			
		# Get bacteries speciments
		for key, d in data.items():
			column_now = 6
			specimentsTemp = {}
			totalValueTemp = 0
			for cell in ws.iter_cols(min_col=6, max_col=wb_column_organisma_end, min_row=d['bactery_row'], max_row=d['bactery_row']):
				for c in cell:
					if c.value:
						for s in specimentColumns:
							if column_now == s['column_position']:
								if s['speciment'] in dictionary_speciments:								
									specimentsTemp[s['speciment']] = {"alias": s['speciment'], "name": dictionary_speciments[s['speciment']]['name'], "value": int(c.value)}
									totalValueTemp += int(c.value)
								else:
									specimentsTemp[s['speciment']] = {"alias": s['speciment'], "name": "unknown", "value": int(c.value)}
				d['speciments'] = specimentsTemp
				column_now += 1
			d['speciments']["total_value"] = totalValueTemp
		return data
		
	def sheetAppendBacteriesWithAntibioticsRow(self, ws = None, bacteries = [], start_row = 0):
		# Update bacteries with antibiotic's row
		data = bacteries
		for wb_row_index in range(start_row, ws.max_row):
			if ws['A' + str(wb_row_index)].value:
				if '_' not in ws['A' + str(wb_row_index)].value.lower() and 'sandi' not in ws['A' + str(wb_row_index)].value.lower():
					if ws['A' + str(wb_row_index)].value.lower() in data:
						if wb_row_index > data[ws['A' + str(wb_row_index)].value.lower()]['bactery_row']:
							data[ws['A' + str(wb_row_index)].value.lower()]['antibiotic_row'] = wb_row_index
		return data
		
	def sheetAppendBacteriesWithAntibioticsData(self, ws = None, bacteries = []):
		# Update bacteries with antibiotic's row
		data = bacteries
		for key, value in data.items():
			# Find Sandi or Nama Antibiotik row
			wb_row_sandi_start = value['antibiotic_row']
			wb_row_antibiotics_start = wb_row_sandi_start
			wb_row_antibiotics_end = wb_row_sandi_start
		
			if value['antibiotic_row'] > 0:
				# Find row contain "sandi" word
				for wb_row_index in range(wb_row_sandi_start, wb_row_sandi_start + 2):
					#if ws['A' + str(wb_row_index)].value or ws['B' + str(wb_row_index)].value:
					if ws['B' + str(wb_row_index)].value:
						if ws['B' + str(wb_row_index)].value.lower() == 'nama antibiotik':
						#if ws['A' + str(wb_row_index)].value.lower() == 'sandi' or ws['B' + str(wb_row_index)].value.lower() == 'nama antibiotik':
							wb_row_sandi_start = wb_row_index
							wb_row_antibiotics_start = wb_row_sandi_start + 1
				
				# Find row contain blank, assume that is the last row for antibiotics
				wb_row_index = wb_row_antibiotics_start
				while True:
					if ws['B' + str(wb_row_index)].value:
						value['antibiotics'][ws['A' + str(wb_row_index)].value.lower()] = {"name": ws['B' + str(wb_row_index)].value.lower(), "alias": ws['A' + str(wb_row_index)].value.lower(), "row": wb_row_index, "r": {"value": float(0), "count": 0}, "i": {"value": float(0), "count": 0}, "s": {"value": float(0), "count": 0}}
						wb_row_antibiotics_end = wb_row_index
						wb_row_index += 1
					else:
						break
		return data
		
	def sheetAppendBacteriesWithAntibioticsRIS(self, ws = None, bacteries = []):
		# Update bacteries with antibiotic's row
		data = bacteries
		for key, value in data.items():
			if value['antibiotic_row'] > 0:
				for k, v in value['antibiotics'].items():
					if k != "r" and k != "i" and k != "s":
						if ws['E' + str(v['row'])].value:
							v['r']['value'] = float(ws['E' + str(v['row'])].value)
							v['r']['count'] = int(ws['D' + str(v['row'])].value)
						if ws['F' + str(v['row'])].value:
							v['i']['value'] = float(ws['F' + str(v['row'])].value)
							v['i']['count'] = int(ws['D' + str(v['row'])].value)
						if ws['G' + str(v['row'])].value:
							v['s']['value'] = float(ws['G' + str(v['row'])].value)
							v['s']['count'] = int(ws['D' + str(v['row'])].value)
						#print(v)
		return data
	
	def sheetAppendBacteriesWithAntibioticsRISMax(self, ws = None, bacteries = []):
		# Update bacteries with antibiotic's row
		data = bacteries
		for key, value in data.items():
			maxR = 0.0
			maxI = 0.0
			maxS = 0.0
			if value['antibiotic_row'] > 0:
				for k, v in value['antibiotics'].items():
					if k != "r" and k != "i" and k != "s":
						if float(v['r']['value']) > maxR:
							maxR = float(v['r']['value'])
						if float(v['i']['value']) > maxI:
							maxI = float(v['i']['value'])
						if float(v['s']['value']) > maxS:
							maxS = float(v['s']['value'])
				value['antibiotics']['r']['max_value'] = maxR
				value['antibiotics']['i']['max_value'] = maxI
				value['antibiotics']['s']['max_value'] = maxS
				# Fill antibiotics that value same as max value
				antibioticsR = {}
				antibioticsI = {}
				antibioticsS = {}
				antibioticsRRow = 0
				antibioticsIRow = 0
				antibioticsSRow = 0
				for k, v in value['antibiotics'].items():
					if k != "r" and k != "i" and k != "s":
						if maxR > 0.0 and float(v['r']['value']) == maxR:
							antibioticsR[v['alias']] = {"alias": v['alias'], "name": v['name'], "count": int(v['r']['count'])}
							antibioticsRRow += 1
						if maxI > 0.0 and float(v['i']['value']) == maxI:
							antibioticsI[v['alias']] = {"alias": v['alias'], "name": v['name'], "count": int(v['i']['count'])}
							antibioticsIRow += 1
						if maxS > 0.0 and float(v['s']['value']) == maxS:
							antibioticsS[v['alias']] = {"alias": v['alias'], "name": v['name'], "count": int(v['s']['count'])}
							antibioticsSRow += 1
				value['antibiotics']['r']['antibiotics'] = antibioticsR
				value['antibiotics']['i']['antibiotics'] = antibioticsI
				value['antibiotics']['s']['antibiotics'] = antibioticsS
				value['antibiotics']['r']['antibiotics_row'] = antibioticsRRow
				value['antibiotics']['i']['antibiotics_row'] = antibioticsIRow
				value['antibiotics']['s']['antibiotics_row'] = antibioticsSRow
		return data
		
	def tableAddRooms(self, bacteries = {}):
		data = {}
		for key, value in bacteries.items():
			#data[key] = {"name": value['name'].title(), "data_exists": 0}
			data[key] = {"name": value['name'], "data_exists": 0}
			if any(value['data']):
				data[key]['data_exists'] = 1
				data[key]['bacteries'] = {"speciments": {}, "bacteries": {}}
				data[key]['r'] = {"antibiotics_row": 0}
				data[key]['i'] = {"antibiotics_row": 0}
				data[key]['s'] = {"antibiotics_row": 0}
		return data
		
	def tableAddBacteries(self, room_alias = 'l1ad', data_room = {}):
		data = {}
		#for key, value in data_room.items():
		datum = data_room['data']['bacteries']
		for key, value in datum.items():
			if int(value['antibiotic_row']) > 0:
				data[key] = {"alias": key, "name": value['name'], "speciments": value['speciments']}
			#print(k)
		#print(room_alias)
		#print(data_room)
		return data
		
	def tableAddSpeciments(self, room_alias = 'l1ad', data_room = {}):
		data = {}
		tempList = []
		datum = data_room['data']['bacteries']
		for key, value in datum.items():
			if int(value['antibiotic_row']) > 0:
				for k, v in value['speciments'].items():
					if k != 'total_value' and k not in tempList:
						tempList.append(k)
						data[k] = {"alias": v['alias'], "name": v['name']}
		return data
		
	def tableAddR(self, room_alias = 'l1ad', data_room = {}):
		data = {}
		tempList = []
		datum = data_room['data']['bacteries']
		for key, value in datum.items():
			if value['antibiotics']['r']['antibiotics_row'] > 0:
				data[key] = {"alias": value['alias'], "name": value['name'], "antibiotics_row": int(value['antibiotics']['r']['antibiotics_row']), "antibiotics_max_value": float(value['antibiotics']['r']['max_value']), "antibiotics": value['antibiotics']['r']['antibiotics']}
		return data
	
	def tableAddI(self, room_alias = 'l1ad', data_room = {}):
		data = {}
		tempList = []
		datum = data_room['data']['bacteries']
		for key, value in datum.items():
			if value['antibiotics']['i']['antibiotics_row'] > 0:
				data[key] = {"alias": value['alias'], "name": value['name'], "antibiotics_row": int(value['antibiotics']['i']['antibiotics_row']), "antibiotics_max_value": float(value['antibiotics']['i']['max_value']), "antibiotics": value['antibiotics']['i']['antibiotics']}
		return data
		
	def tableAddS(self, room_alias = 'l1ad', data_room = {}):
		data = {}
		tempList = []
		datum = data_room['data']['bacteries']
		for key, value in datum.items():
			if value['antibiotics']['s']['antibiotics_row'] > 0:
				data[key] = {"alias": value['alias'], "name": value['name'], "antibiotics_row": int(value['antibiotics']['s']['antibiotics_row']), "antibiotics_max_value": float(value['antibiotics']['s']['max_value']), "antibiotics": value['antibiotics']['s']['antibiotics']}
		return data
			
	def addChart(self, table = {}, month = 5, year = 2017, months = None):
		import PIL
		import matplotlib.pyplot as plt; plt.rcdefaults()
		import numpy as np
		import matplotlib.pyplot as plt
		from os import path
		from os import makedirs
		from PIL import Image
		
		monthName = 'mei'
		if str(month) in months:
			monthName = months[str(month)]['name']
		
		newDir = monthName + '_' + str(year)
		if not path.exists(newDir):
			makedirs(newDir)
			
		# Foreach rooms
		for key, value in table.items():
			print('Create chart for room:', value['name'])
			fileName = monthName + '_' + str(year) + '_'
			if int(value['data_exists']) == 1:
				title = []
				number = []
				
				for k, v in value['bacteries']['bacteries'].items():
					title.append(k)
					number.append(v['speciments']['total_value'])				
				
				fig = plt.figure()
				ax = plt.subplot(111)
				width = 0.8
				ax.bar(range(len(title)), number, width=width, align='center')
				ax.set_xticks(np.arange(len(title)) + width/2)
				ax.set_xticklabels(title)
				ax.set_title(value['name'].title(), ha='center')
				#ax.set_xticklabels(title, rotation=90)
				fileName = fileName + key + '.png'
				fileName = newDir + '/' + fileName
				#plt.savefig(fileName)
				fig.savefig(fileName)
				plt.close(fig)
				
				# Resize for Excel
				baseWidth = 480
				img = Image.open(fileName)
				wPercent = (baseWidth/float(img.size[0]))
				hSize = int((float(img.size[1]) * float(wPercent)))
				img = img.resize((baseWidth, hSize), PIL.Image.ANTIALIAS)
				fileName = monthName + '_' + str(year) + '_'
				fileName = fileName + '480_' + key + '.png'
				fileName = newDir + '/' + fileName
				img.save(fileName)
				
				# Resize for Word
				baseWidth = 360
				img = Image.open(fileName)
				wPercent = (baseWidth/float(img.size[0]))
				hSize = int((float(img.size[1]) * float(wPercent)))
				img = img.resize((baseWidth, hSize), PIL.Image.ANTIALIAS)
				fileName = monthName + '_' + str(year) + '_'
				fileName = fileName + '360_' + key + '.png'
				fileName = newDir + '/' + fileName
				img.save(fileName)
	
	def newDoc(self, data = {}, month = 5, year = 2017, months = None, rooms = {}, max_columns = 8):
		from os import path
		from os import makedirs
		from docx import Document
		from docx.shared import Inches
		from docx.enum.text import WD_ALIGN_PARAGRAPH
		from docx.enum.section import WD_ORIENT
		from docx.enum.table import WD_TABLE_ALIGNMENT
		#from docx.oxml.ns import nsdecls
		#from docx.oxml import parse_xml

		# Set a cell background (shading) color to RGB D9D9D9. 
		#shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
		#cell._tc.get_or_add_tcPr().append(shading_elm)
		
		print('Creating docx...')
		#columnSheet = list(string.ascii_uppercase)
		
		# Create new file
		monthName = 'mei'
		if str(month) in months:
			monthName = months[str(month)]['name']
		
		newDir = monthName + '_' + str(year)
		if not path.exists(newDir):
			makedirs(newDir)
		
		fileName = newDir + '/' + monthName + '_' + str(year) + '.docx'
		
		document = Document()
		
		# Page Setup
		sections = document.sections
		section = sections[0]
		section.page_width = Inches(11.69)
		section.page_height = Inches(8.27)
		section.left_margin = Inches(1)
		section.right_margin = Inches(1)
		section.top_margin = Inches(1)
		section.bottom_margin = Inches(1)
		section.orientation = WD_ORIENT.LANDSCAPE
		
		# Create Title
		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER		
		p.add_run(monthName.upper() + ' ' + str(year)).bold = True

		# Iterate over rooms
		# Foreach rooms
		roomNumber = 1
		for key, value in rooms.items():
			print('Writing ' + value['name'])
			# Room title
			p = document.add_paragraph()
			p = document.add_paragraph()
			p.add_run(str(roomNumber) + '. ' + value['name'].upper()).bold = True
			roomNumber += 1
			
			# Room table
			if key in data:
				if int(data[key]['data_exists']) == 1:
					p = document.add_paragraph('Tabel 1. Distribusi bakteri terbanyak menurut jenis spesimen')
					
					table = document.add_table(rows=1, cols=(len(data[key]['bacteries']['speciments'])) + 3)
					# Bacterie's table header
					table.cell(0, 0).text = 'No'
					table.cell(0, 0).width = Inches(0.4)
					table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					table.cell(0, 1).text = 'Bakteri/Spesimen'
					table.cell(0, 1).width = Inches(1.6)
					table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					# Print speciments
					columnCounter = 2
					for k, v in data[key]['bacteries']['speciments'].items():
						table.cell(0, columnCounter).text = v['name'].title()
						table.cell(0, columnCounter).width = Inches(0.8)
						table.cell(0, columnCounter).paragraphs[0].runs[0].font.bold = True
						table.cell(0, columnCounter).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
						# Update table's data
						v['column_index'] = columnCounter
						columnCounter += 1
					# Total
					table.cell(0, columnCounter).text = 'Total'
					table.cell(0, columnCounter).width = Inches(0.8)
					table.cell(0, columnCounter).paragraphs[0].runs[0].font.bold = True
					table.cell(0, columnCounter).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					# Bacterie's bacteries data
					# Print bacteries
					number = 1
					rowCounter = 1
					for k, v in data[key]['bacteries']['bacteries'].items():
						table.add_row()
						
						table.cell(rowCounter, 0).text = str(number)
						table.cell(rowCounter, 0).width = Inches(0.4)
						table.cell(rowCounter, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
						
						table.cell(rowCounter, 1).text = v['name'].title()
						table.cell(rowCounter, 1).width = Inches(1.6)
						table.cell(rowCounter, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
						
						number += 1
						
						# Bactery's Speciments
						for a, b in v['speciments'].items():
							if a != 'total_value':
								speciments = data[key]['bacteries']['speciments']
								
								table.cell(rowCounter, int(speciments[a]['column_index'])).text = str(b['value'])
								table.cell(rowCounter, int(speciments[a]['column_index'])).width = Inches(0.8)
								table.cell(rowCounter, int(speciments[a]['column_index'])).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
								# Update table's data
								b['column_index'] = speciments[a]['column_index']
								
						# Total column						
						table.cell(rowCounter, columnCounter).text = str(v['speciments']['total_value'])
						table.cell(rowCounter, columnCounter).width = Inches(0.8)
						table.cell(rowCounter, columnCounter).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
						
						rowCounter += 1
					table.style = 'Table Grid'
					table.alignment = WD_TABLE_ALIGNMENT.LEFT
					
					# Insert chart image
					p = document.add_paragraph()
					p = document.add_paragraph()
					p.add_run().add_picture(newDir + '/' + monthName + '_' + str(year) + '_360_' + key + '.png')
					
					# Print R Table, if data exists
					# R table get maximum column
					tableMaxColumn = 0
					if any(data[key]['r']):
						for k, v in data[key]['r'].items():
							if v['antibiotics_row'] > tableMaxColumn:
								tableMaxColumn = v['antibiotics_row']
					if int(tableMaxColumn) > int(max_columns):
						tableMaxColumn = int(max_columns)
					
					# R table header
					p = document.add_paragraph()
					if tableMaxColumn == 0:
						table = document.add_table(rows=1, cols=3)
					else:
						table = document.add_table(rows=1, cols=(tableMaxColumn + 2))
					
					# Bactery's table header
					table.cell(0, 0).text = 'No'
					table.cell(0, 0).width = Inches(0.4)
					table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					table.cell(0, 1).text = 'Nama Kuman'
					table.cell(0, 1).width = Inches(1.2)
					table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					table.cell(0, 2).text = 'Antibiotik Resisten (%R)'
					table.cell(0, 2).width = Inches(1.1)
					table.cell(0, 2).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					if tableMaxColumn > 0:
						a = table.cell(0, 2)
						b = table.cell(0, (tableMaxColumn + 1))
						a.merge(b)
					
					# If R table contain with data
					if any(data[key]['r']):
						tableMaxColumn = 0
						number = 1
						rowCounter = 1
						for k, v in data[key]['r'].items():
							antibiotics = v['antibiotics']
							table.add_row()
							table.add_row()
							# Print bacteries name
							table.cell(rowCounter, 0).text = str(number)
							table.cell(rowCounter, 0).width = Inches(0.4)
							table.cell(rowCounter, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
							
							number += 1
							table.cell(rowCounter, 1).text = v['name'].title()
							table.cell(rowCounter, 1).width = Inches(1.2)
							#table.cell(rowCounter, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
							
							# Print bacteries values
							columnCounter = 2
							columnCount = 0
							for a, b in antibiotics.items():
								titleRow = rowCounter
								valueRow = titleRow + 1								
								table.cell(titleRow, columnCounter).text = b['name'].title()
								table.cell(titleRow, columnCounter).width = Inches(1.1)
								table.cell(titleRow, columnCounter).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
								
								val = v['antibiotics_max_value']
								if float(val)%10 == 0.0:
									val = int(val)
								
								table.cell(valueRow, columnCounter).text = str(b['count']) + '(' + str(val) + ')'
								table.cell(valueRow, columnCounter).width = Inches(1.1)
								table.cell(valueRow, columnCounter).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
								
								columnCounter += 1
								columnCount += 1
								if columnCount >= int(max_columns) and int(v['antibiotics_row']) > int(max_columns):
									columnCount = 0
									columnCounter = 2
									rowCounter += 2
									table.add_row()
									table.add_row()
							rowCounter += 2
									
					# If no data to show on table R
					else:
						table.add_row()
					table.style = 'Table Grid'
					table.alignment = WD_TABLE_ALIGNMENT.LEFT
					
					# Print S Table, if data exists
					# S table get maximum column
					tableMaxColumn = 0
					if any(data[key]['s']):
						for k, v in data[key]['s'].items():
							if v['antibiotics_row'] > tableMaxColumn:
								tableMaxColumn = v['antibiotics_row']
					if int(tableMaxColumn) > int(max_columns):
						tableMaxColumn = int(max_columns)
					
					# S table header
					p = document.add_paragraph()
					if tableMaxColumn == 0:
						table = document.add_table(rows=1, cols=3)
					else:
						table = document.add_table(rows=1, cols=(tableMaxColumn + 2))
					
					# Bactery's table header
					table.cell(0, 0).text = 'No'
					table.cell(0, 0).width = Inches(0.4)
					table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					table.cell(0, 1).text = 'Nama Kuman'
					table.cell(0, 1).width = Inches(1.2)
					table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					table.cell(0, 2).text = 'Antibiotik Sensitive (%S)'
					table.cell(0, 2).width = Inches(1.1)
					table.cell(0, 2).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					if tableMaxColumn > 0:
						a = table.cell(0, 2)
						b = table.cell(0, (tableMaxColumn + 1))
						a.merge(b)
					
					# If S table contain with data
					if any(data[key]['s']):
						tableMaxColumn = 0
						number = 1
						rowCounter = 1
						for k, v in data[key]['s'].items():
							antibiotics = v['antibiotics']
							table.add_row()
							table.add_row()
							# Print bacteries name
							table.cell(rowCounter, 0).text = str(number)
							table.cell(rowCounter, 0).width = Inches(0.4)
							table.cell(rowCounter, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
							
							number += 1
							table.cell(rowCounter, 1).text = v['name'].title()
							table.cell(rowCounter, 1).width = Inches(1.2)
							#table.cell(rowCounter, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
							
							# Print bacteries values
							columnCounter = 2
							columnCount = 0
							for a, b in antibiotics.items():
								titleRow = rowCounter
								valueRow = titleRow + 1								
								table.cell(titleRow, columnCounter).text = b['name'].title()
								table.cell(titleRow, columnCounter).width = Inches(1.1)
								table.cell(titleRow, columnCounter).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
								
								val = v['antibiotics_max_value']
								if float(val)%10 == 0.0:
									val = int(val)
								
								table.cell(valueRow, columnCounter).text = str(b['count']) + '(' + str(val) + ')'
								table.cell(valueRow, columnCounter).width = Inches(1.1)
								table.cell(valueRow, columnCounter).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
								
								columnCounter += 1
								columnCount += 1
								if columnCount >= int(max_columns) and int(v['antibiotics_row']) > int(max_columns):
									columnCount = 0
									columnCounter = 2
									rowCounter += 2
									table.add_row()
									table.add_row()
							rowCounter += 2
									
					# If no data to show on table S
					else:
						table.add_row()
					table.style = 'Table Grid'
					table.alignment = WD_TABLE_ALIGNMENT.LEFT
					
					# Print I Table, if data exists
					# I table get maximum column
					tableMaxColumn = 0
					if any(data[key]['i']):
						for k, v in data[key]['i'].items():
							if v['antibiotics_row'] > tableMaxColumn:
								tableMaxColumn = v['antibiotics_row']
					if int(tableMaxColumn) > int(max_columns):
						tableMaxColumn = int(max_columns)
					
					# I table header
					p = document.add_paragraph()
					if tableMaxColumn == 0:
						table = document.add_table(rows=1, cols=3)
					else:
						table = document.add_table(rows=1, cols=(tableMaxColumn + 2))
					
					# Bactery's table header
					table.cell(0, 0).text = 'No'
					table.cell(0, 0).width = Inches(0.4)
					table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					table.cell(0, 1).text = 'Nama Kuman'
					table.cell(0, 1).width = Inches(1.2)
					table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					table.cell(0, 2).text = 'Antibiotik Intermediate (%I)'
					table.cell(0, 2).width = Inches(1.1)
					table.cell(0, 2).paragraphs[0].runs[0].font.bold = True
					table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
					
					if tableMaxColumn > 0:
						a = table.cell(0, 2)
						b = table.cell(0, (tableMaxColumn + 1))
						a.merge(b)
					
					# If R table contain with data
					if any(data[key]['i']):
						tableMaxColumn = 0
						number = 1
						rowCounter = 1
						for k, v in data[key]['i'].items():
							antibiotics = v['antibiotics']
							table.add_row()
							table.add_row()
							# Print bacteries name
							table.cell(rowCounter, 0).text = str(number)
							table.cell(rowCounter, 0).width = Inches(0.4)
							table.cell(rowCounter, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
							
							number += 1
							table.cell(rowCounter, 1).text = v['name'].title()
							table.cell(rowCounter, 1).width = Inches(1.2)
							#table.cell(rowCounter, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
							
							# Print bacteries values
							columnCounter = 2
							columnCount = 0
							for a, b in antibiotics.items():
								titleRow = rowCounter
								valueRow = titleRow + 1								
								table.cell(titleRow, columnCounter).text = b['name'].title()
								table.cell(titleRow, columnCounter).width = Inches(1.1)
								table.cell(titleRow, columnCounter).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
								
								val = v['antibiotics_max_value']
								if float(val)%10 == 0.0:
									val = int(val)
								
								table.cell(valueRow, columnCounter).text = str(b['count']) + '(' + str(val) + ')'
								table.cell(valueRow, columnCounter).width = Inches(1.1)
								table.cell(valueRow, columnCounter).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
								
								columnCounter += 1
								columnCount += 1
								if columnCount >= int(max_columns) and int(v['antibiotics_row']) > int(max_columns):
									columnCount = 0
									columnCounter = 2
									rowCounter += 2
									table.add_row()
									table.add_row()
							rowCounter += 2
									
					# If no data to show on table I
					else:
						table.add_row()
					table.style = 'Table Grid'
					table.alignment = WD_TABLE_ALIGNMENT.LEFT
					
				# If no data to show
				else:
					p = document.add_paragraph('TIDAK ADA')
			# If no data to show
			else:
				p = document.add_paragraph('TIDAK ADA')
		# Save document
		document.save(fileName)
		return data
	
	def newFile(self, table = {}, month = 5, year = 2017, months = None, rooms = {}, max_columns = 8):
		from os import path
		from os import makedirs
		from openpyxl.drawing.image import Image
		
		print('Creating xlsx...')
		
		# Buat file baru
		monthName = 'mei'
		if str(month) in months:
			monthName = months[str(month)]['name']
		
		newDir = monthName + '_' + str(year)
		if not path.exists(newDir):
			makedirs(newDir)
		
		fileName = newDir + '/' + monthName + '_' + str(year) + '.xlsx'
		wb = Workbook()
		# Worksheet's name
		ws = wb.create_sheet(title=monthName.title())
		#styleBold = Style(font=Font(bold=True))
		
		# Set column width
		ws.column_dimensions['A'].width = 4
		ws.column_dimensions['B'].width = 17
		ws.column_dimensions['C'].width = 12
		ws.column_dimensions['D'].width = 12
		ws.column_dimensions['E'].width = 12
		ws.column_dimensions['F'].width = 12
		ws.column_dimensions['G'].width = 12
		ws.column_dimensions['H'].width = 12
		ws.column_dimensions['I'].width = 12
		ws.column_dimensions['J'].width = 12
		
		# Styles
		Side = styles.Side
		sBold = styles.Font(bold=True)
		sCenter = styles.Alignment(horizontal='center')
		sCenterWrap = styles.Alignment(horizontal='center', vertical='top', wrap_text=True)
		sWrap = styles.Alignment(vertical='top', wrap_text=True)
		sBorder = styles.Border(left=Side(border_style='thin'), right=Side(border_style='thin'), top=Side(border_style='thin'), bottom=Side(border_style='thin'))
		#sRed = styles.PatternFill(bgColor='FF0000', fill_type='solid')
		#sGreen = styles.PatternFill(bgColor='15FF00', fill_type='solid')
		#sYellow = styles.PatternFill(bgColor='FFF600', fill_type='solid')
		sRed = styles.PatternFill('solid', styles.colors.RED)
		sGreen = styles.PatternFill('solid', styles.colors.GREEN)
		sYellow = styles.PatternFill('solid', styles.colors.YELLOW)
		
		# Create Title
		ws['A1'] = monthName.title() + ' ' + str(year)
		_cell = ws['A1']
		_cell.font = sBold
		_cell.alignment = sCenter
		ws.merge_cells('A1:J1')
		
		actualRow = 3
		
		# Daftar column
		#columnSheet = list(string.ascii_uppercase)
		#columnSheetIndex = 2 # start from C column
		# columnSheet = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
		
		# Iterate over rooms
		# Foreach rooms
		for key, value in rooms.items():
			print('Writing ' + value['name'])
			# Room title
			ws['A' + str(actualRow)] = value['name'].upper()
			_cell = ws['A' + str(actualRow)]
			_cell.font = sBold
			actualRow += 1
			
			# Room table
			if key in table:
				if int(table[key]['data_exists']) == 1:
					ws['A' + str(actualRow)] = 'Tabel 1. Distribusi bakteri terbanyak menurut jenis spesimen'
					_cell = ws['A' + str(actualRow)]
					_cell.font = sBold
					
					# Bacterie's table header
					actualRow += 1
					ws['A' + str(actualRow)] = 'No'
					_cell = ws['A' + str(actualRow)]
					_cell.font = sBold
					_cell.alignment = sCenterWrap
					_cell.border = sBorder
					ws['B' + str(actualRow)] = 'Bakteri/Spesimen'
					_cell = ws['B' + str(actualRow)]
					_cell.font = sBold
					_cell.alignment = sCenterWrap
					_cell.border = sBorder
					
					columnSheet = list(string.ascii_uppercase)
					columnSheetIndex = 2 # start from C column
					# columnSheet = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
					# Print speciment
					for k, v in table[key]['bacteries']['speciments'].items():
						ws[columnSheet[columnSheetIndex] + str(actualRow)] = v['name'].title()
						_cell = ws[columnSheet[columnSheetIndex] + str(actualRow)]
						_cell.font = sBold
						_cell.alignment = sCenterWrap
						_cell.border = sBorder
						
						# Update table's data
						v['column'] = columnSheet[columnSheetIndex]
						columnSheetIndex += 1
					# Total column
					ws[columnSheet[columnSheetIndex] + str(actualRow)] = 'Total'
					_cell = ws[columnSheet[columnSheetIndex] + str(actualRow)]
					_cell.font = sBold
					_cell.alignment = sCenterWrap
					_cell.border = sBorder
					
					# Bacterie's bacteries data
					# Print bacteries
					actualRow += 1
					number = 1
					for k, v in table[key]['bacteries']['bacteries'].items():
						ws['A' + str(actualRow)] = number
						_cell = ws['A' + str(actualRow)]
						_cell.alignment = sCenterWrap
						_cell.border = sBorder
						ws['B' + str(actualRow)] = v['name'].title()
						_cell = ws['B' + str(actualRow)]
						_cell.alignment = sWrap
						_cell.border = sBorder
						
						# Bactery's Speciments
						for a, b in v['speciments'].items():
							if a != 'total_value':
								speciments = table[key]['bacteries']['speciments']
								ws[str(speciments[a]['column']) + str(actualRow)] = b['value']
								_cell = ws[str(speciments[a]['column']) + str(actualRow)]
								_cell.alignment = sCenterWrap
								#_cell.border = sBorder
								
								# Update table's data
								b['column'] = speciments[a]['column']
								
						# Total column
						ws[columnSheet[columnSheetIndex] + str(actualRow)] = '=SUM(C' + str(actualRow) + ':' + columnSheet[columnSheetIndex-1] + str(actualRow) + ')'
						_cell = ws[columnSheet[columnSheetIndex] + str(actualRow)]
						_cell.alignment = sCenterWrap
						_cell.border = sBorder
						
						# Create border
						for c in range(2, columnSheetIndex):
							_cell = ws[columnSheet[c] + str(actualRow)]
							_cell.border = sBorder
							
						number += 1
						actualRow += 1
					
					# Insert Chart's Image
					actualRow += 2
					img = Image(newDir + '/' + monthName + '_' + str(year) + '_480_' + key + '.png')
					imgColumn = 'A' + str(actualRow)
					ws.add_image(img, imgColumn)
					
					actualRow += 20
					
					# Print R Table
					# R table header
					ws['A' + str(actualRow)] = 'No'
					_cell = ws['A' + str(actualRow)]
					_cell.font = sBold
					_cell.alignment = sCenterWrap
					_cell.border = sBorder
					ws['B' + str(actualRow)] = 'Nama Kuman'
					_cell = ws['B' + str(actualRow)]
					_cell.font = sBold
					_cell.alignment = sCenterWrap
					_cell.border = sBorder
					# If R table contain with data
					if any(table[key]['r']):
						tableHeaderRow = actualRow
						tableMaxColumn = 0
						number = 1
						for k, v in table[key]['r'].items():
							if v['antibiotics_row'] > tableMaxColumn:
								tableMaxColumn = v['antibiotics_row']
							columnCount = 0
							antibiotics = v['antibiotics']
							actualRow += 1
							# Print bacteries name
							ws['A' + str(actualRow)] = number
							number += 1
							_cell = ws['A' + str(actualRow)]
							_cell.alignment = sCenterWrap
							_cell.border = sBorder
							ws['B' + str(actualRow)] = v['name'].title()
							_cell = ws['B' + str(actualRow)]
							_cell.alignment = sWrap
							_cell.border = sBorder
							
							# Print bacteries values
							for a, b in antibiotics.items():
								titleRow = actualRow
								valueRow = titleRow + 1
								ws[columnSheet[columnCount + 2] + str(titleRow)] = b['name'].title()
								_cell = ws[columnSheet[columnCount + 2] + str(titleRow)]
								_cell.alignment = sCenterWrap
								_cell.border = sBorder
								
								val = v['antibiotics_max_value']
								if float(val)%10 == 0.0:
									val = int(val)
									
								ws[columnSheet[columnCount + 2] + str(valueRow)] = str(b['count']) + '(' + str(val) + ')'
								_cell = ws[columnSheet[columnCount + 2] + str(valueRow)]
								_cell.alignment = sCenter
								_cell.border = sBorder
								_cell.fill = sRed
								columnCount += 1
								#if columnCount >= int(max_columns):
								if columnCount >= int(max_columns) and int(v['antibiotics_row']) > int(max_columns):
									columnCount = 0
									actualRow += 2
							actualRow += 1
						
						# Print Table Header
						if int(tableMaxColumn) > int(max_columns):
							tableMaxColumn = int(max_columns)
						
						# Create table title header
						ws['C' + str(tableHeaderRow)] = 'Antibiotik Resisten (%R)'
						_cell = ws['C' + str(tableHeaderRow)]
						_cell.font = sBold
						_cell.alignment = sCenterWrap
						
						# Last column
						lastColumn = columnSheet[tableMaxColumn + 1]
						
						ws.merge_cells('C' + str(tableHeaderRow) + ':' + lastColumn + str(tableHeaderRow))
						
						# Create Border
						actualRow += 1
						for c in range(tableHeaderRow, actualRow):
							for d in range(0, (tableMaxColumn + 2)):
								_cell = ws[columnSheet[d] + str(c)]
								_cell.border = sBorder
					
					# If R table is empty
					else:
						ws['C' + str(actualRow)] = 'Antibiotik Resisten (%R)'
						_cell = ws['C' + str(actualRow)]
						_cell.font = sBold
						_cell.alignment = sCenterWrap
						_cell.border = sBorder
						actualRow += 1
						_cell = ws['A' + str(actualRow)]
						_cell.border = sBorder
						_cell = ws['B' + str(actualRow)]
						_cell.border = sBorder
						_cell = ws['C' + str(actualRow)]
						_cell.border = sBorder
					actualRow += 1
					
					# Print I Table
					# I table header
					ws['A' + str(actualRow)] = 'No'
					_cell = ws['A' + str(actualRow)]
					_cell.font = sBold
					_cell.alignment = sCenterWrap
					_cell.border = sBorder
					ws['B' + str(actualRow)] = 'Nama Kuman'
					_cell = ws['B' + str(actualRow)]
					_cell.font = sBold
					_cell.alignment = sCenterWrap
					_cell.border = sBorder
					# If I table contain with data
					if any(table[key]['i']):
						tableHeaderRow = actualRow
						tableMaxColumn = 0
						number = 1
						for k, v in table[key]['i'].items():
							if v['antibiotics_row'] > tableMaxColumn:
								tableMaxColumn = v['antibiotics_row']
							columnCount = 0
							antibiotics = v['antibiotics']
							actualRow += 1
							# Print bacteries name
							ws['A' + str(actualRow)] = number
							number += 1
							_cell = ws['A' + str(actualRow)]
							_cell.alignment = sCenterWrap
							_cell.border = sBorder
							ws['B' + str(actualRow)] = v['name'].title()
							_cell = ws['B' + str(actualRow)]
							_cell.alignment = sWrap
							_cell.border = sBorder
							
							# Print bacteries values
							for a, b in antibiotics.items():
								titleRow = actualRow
								valueRow = titleRow + 1
								ws[columnSheet[columnCount + 2] + str(titleRow)] = b['name'].title()
								_cell = ws[columnSheet[columnCount + 2] + str(titleRow)]
								_cell.alignment = sCenterWrap
								_cell.border = sBorder
								
								val = v['antibiotics_max_value']
								if float(val)%10 == 0.0:
									val = int(val)
									
								ws[columnSheet[columnCount + 2] + str(valueRow)] = str(b['count']) + '(' + str(val) + ')'
								_cell = ws[columnSheet[columnCount + 2] + str(valueRow)]
								_cell.alignment = sCenter
								_cell.border = sBorder
								_cell.fill = sGreen
								columnCount += 1
								#if columnCount >= int(max_columns):
								if columnCount >= int(max_columns) and int(v['antibiotics_row']) > int(max_columns):
									columnCount = 0
									actualRow += 2
							actualRow += 1
						
						# Print Table Header
						if int(tableMaxColumn) > int(max_columns):
							tableMaxColumn = int(max_columns)
						
						# Create table title header
						ws['C' + str(tableHeaderRow)] = 'Antibiotik Intermediate (%I)'
						_cell = ws['C' + str(tableHeaderRow)]
						_cell.font = sBold
						_cell.alignment = sCenterWrap
						
						# Last column
						lastColumn = columnSheet[tableMaxColumn + 1]
						
						ws.merge_cells('C' + str(tableHeaderRow) + ':' + lastColumn + str(tableHeaderRow))
						
						# Create Border
						actualRow += 1
						for c in range(tableHeaderRow, actualRow):
							for d in range(0, (tableMaxColumn + 2)):
								_cell = ws[columnSheet[d] + str(c)]
								_cell.border = sBorder
					
					# If I table is empty
					else:
						ws['C' + str(actualRow)] = 'Antibiotik Intermediate (%I)'
						_cell = ws['C' + str(actualRow)]
						_cell.font = sBold
						_cell.alignment = sCenterWrap
						_cell.border = sBorder
						actualRow += 1
						_cell = ws['A' + str(actualRow)]
						_cell.border = sBorder
						_cell = ws['B' + str(actualRow)]
						_cell.border = sBorder
						_cell = ws['C' + str(actualRow)]
						_cell.border = sBorder
					actualRow += 1
					
					# Print S Table
					# S table header
					ws['A' + str(actualRow)] = 'No'
					_cell = ws['A' + str(actualRow)]
					_cell.font = sBold
					_cell.alignment = sCenterWrap
					_cell.border = sBorder
					ws['B' + str(actualRow)] = 'Nama Kuman'
					_cell = ws['B' + str(actualRow)]
					_cell.font = sBold
					_cell.alignment = sCenterWrap
					_cell.border = sBorder
					# If S table contain with data
					if any(table[key]['s']):
						tableHeaderRow = actualRow
						tableMaxColumn = 0
						number = 1
						for k, v in table[key]['s'].items():
							if v['antibiotics_row'] > tableMaxColumn:
								tableMaxColumn = v['antibiotics_row']
							columnCount = 0
							antibiotics = v['antibiotics']
							actualRow += 1
							# Print bacteries name
							ws['A' + str(actualRow)] = number
							number += 1
							_cell = ws['A' + str(actualRow)]
							_cell.alignment = sCenterWrap
							_cell.border = sBorder
							ws['B' + str(actualRow)] = v['name'].title()
							_cell = ws['B' + str(actualRow)]
							_cell.alignment = sWrap
							_cell.border = sBorder
							
							# Print bacteries values
							for a, b in antibiotics.items():
								titleRow = actualRow
								valueRow = titleRow + 1
								ws[columnSheet[columnCount + 2] + str(titleRow)] = b['name'].title()
								_cell = ws[columnSheet[columnCount + 2] + str(titleRow)]
								_cell.alignment = sCenterWrap
								_cell.border = sBorder
								
								val = v['antibiotics_max_value']
								if float(val)%10 == 0.0:
									val = int(val)
									
								ws[columnSheet[columnCount + 2] + str(valueRow)] = str(b['count']) + '(' + str(val) + ')'
								_cell = ws[columnSheet[columnCount + 2] + str(valueRow)]
								_cell.alignment = sCenter
								_cell.border = sBorder
								_cell.fill = sYellow
								columnCount += 1
								#if columnCount >= int(max_columns):
								if columnCount >= int(max_columns) and int(v['antibiotics_row']) > int(max_columns):
									columnCount = 0
									actualRow += 2
							actualRow += 1
						
						# Print Table Header
						if int(tableMaxColumn) > int(max_columns):
							tableMaxColumn = int(max_columns)
						
						# Create table title header
						ws['C' + str(tableHeaderRow)] = 'Antibiotik Sensitive (%S)'
						_cell = ws['C' + str(tableHeaderRow)]
						_cell.font = sBold
						_cell.alignment = sCenterWrap
						
						# Last column
						lastColumn = columnSheet[tableMaxColumn + 1]
						
						ws.merge_cells('C' + str(tableHeaderRow) + ':' + lastColumn + str(tableHeaderRow))
						
						# Create Border
						actualRow += 1
						for c in range(tableHeaderRow, actualRow):
							for d in range(0, (tableMaxColumn + 2)):
								_cell = ws[columnSheet[d] + str(c)]
								_cell.border = sBorder
					
					# If I table is empty
					else:
						ws['C' + str(actualRow)] = 'Antibiotik Sensitive (%S)'
						_cell = ws['C' + str(actualRow)]
						_cell.font = sBold
						_cell.alignment = sCenterWrap
						_cell.border = sBorder
						actualRow += 1
						_cell = ws['A' + str(actualRow)]
						_cell.border = sBorder
						_cell = ws['B' + str(actualRow)]
						_cell.border = sBorder
						_cell = ws['C' + str(actualRow)]
						_cell.border = sBorder
					actualRow += 1
				else:
					ws['A' + str(actualRow)] = 'TIDAK ADA'
					actualRow += 1
			else:
				ws['A' + str(actualRow)] = 'TIDAK ADA'
				actualRow += 1
			actualRow += 1
		
		wb.save(filename=fileName)
		return table
	
	def dumpToJson(self, data = {}, month = 5, year = 2017, months = None, table = {}):
		import json
		from os import path
		from os import makedirs
		
		monthName = 'mei'
		if str(month) in months:
			monthName = months[str(month)]['name']
			
		newDir = monthName + '_' + str(year)
		if not path.exists(newDir):
			makedirs(newDir)
			
		fileName = newDir + '/' + monthName + '_' + str(year) + '.json'
		data = {"month": monthName, "year": year, "data": data, "table": table}
		with open(fileName, 'w') as f:
			json.dump(data, f, indent=4)
		
	def main(self):
		dictionary = self.loadDictionary()
		files = self.showFilesOnDir()
		file = self.pickAFile(files)
		monthYear = self.pickMonthYear(dictionary['settings']['months_long_id'])
		wb = self.openFile(file=file['name'])
		sheets = self.getSheets(wb)
		data = {}
		table = {}
		if len(sheets) > 0:
			for s in sheets:
				# Get room's data
				room = {"name": "unknown", "alias": "unknown", "sheet": "unknown"}
				if s.lower() in dictionary['rooms']:
					room = dictionary['rooms'][s.lower()]
				
				# If sheet contain no data
				sheetNotBlank = self.sheetCheckIfNotBlank(wb[s])
				if sheetNotBlank:					
					print(room['name'], "Ada Data")
					bacteries = self.sheetGetBacteries(wb[s])					
					bacteries['data'] = self.sheetAppendBacteriesWithSpeciments(ws=wb[s], bacteries=bacteries['data'], start_row=bacteries['start_row'], end_row=bacteries['end_row'], organisma_row=bacteries['organisma_row'], dictionary_speciments = dictionary['speciments'])
					#print(room['name'], bacteries['data'])
					#bacteries['data'] = self.sheetAppendBacteriesWithAntibioticsRow(ws=wb[s], bacteries=bacteries['data'], start_row=bacteries['end_row']+1, dictionary_antibiotics=dictionary['antibiotics'])
					#print('bacteries_end_row', str(bacteries['end_row']+1))
					bacteries['data'] = self.sheetAppendBacteriesWithAntibioticsRow(ws=wb[s], bacteries=bacteries['data'], start_row=(bacteries['end_row']+1))
					bacteries['data'] = self.sheetAppendBacteriesWithAntibioticsData(ws=wb[s], bacteries=bacteries['data'])
					bacteries['data'] = self.sheetAppendBacteriesWithAntibioticsRIS(ws=wb[s], bacteries=bacteries['data'])
					bacteries['data'] = self.sheetAppendBacteriesWithAntibioticsRISMax(ws=wb[s], bacteries=bacteries['data'])
					data[room['alias']] = {"name": room['name'], "data": {"bacteries": bacteries['data']}}
				else:
					print(room['name'], "Tidak Ada Data")
					data[room['alias']] = {"name": room['name'], "data": {}}
					
			# Add Table dictionary for a simpler data
			tableRooms = self.tableAddRooms(bacteries=data)
			if any(tableRooms):
				for key, value in tableRooms.items():
					if int(value['data_exists']) == 1:
						tableBacteries = self.tableAddBacteries(room_alias=key, data_room=data[key])
						value['bacteries']['bacteries'] = tableBacteries
						tableSpeciments = self.tableAddSpeciments(room_alias=key, data_room=data[key])
						value['bacteries']['speciments'] = tableSpeciments
						tableR = self.tableAddR(room_alias=key, data_room=data[key])
						value['r'] = tableR
						tableI = self.tableAddI(room_alias=key, data_room=data[key])
						value['i'] = tableI
						tableS = self.tableAddS(room_alias=key, data_room=data[key])
						value['s'] = tableS
			
			table = tableRooms
			#self.addChart(table=table, months=dictionary['settings']['months_long_id'], month=monthYear['month'], year=monthYear['year'])
			table = self.newFile(table=table, months=dictionary['settings']['months_long_id'], month=monthYear['month'], year=monthYear['year'], rooms=dictionary['rooms'], max_columns=dictionary['settings']['columns'])
			table = self.newDoc(data=table, months=dictionary['settings']['months_long_id'], month=monthYear['month'], year=monthYear['year'], rooms=dictionary['rooms'], max_columns=dictionary['settings']['columns'])
			#self.dumpToJson(data=data, months=dictionary['settings']['months_long_id'], month=monthYear['month'], year=monthYear['year'], table=table)
		#print(sheets)
		
def main():
    KumKum()

if __name__ == '__main__':
    main()
